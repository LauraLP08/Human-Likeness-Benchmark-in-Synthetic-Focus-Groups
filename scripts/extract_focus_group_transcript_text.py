import os
import argparse
import json
import shutil

try:
    import fitz
    HAVE_FITZ = True
except ImportError:
    HAVE_FITZ = False

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

def extract_pdf(filepath: str) -> tuple[str, dict]:
    meta = {
        "file_type": "pdf",
        "extraction_method": "pymupdf",
        "pages": 0,
        "character_count": 0,
        "suspected_scanned_pdf": False,
        "ocr_required": False,
        "legacy_doc_conversion_required": False,
        "table_structure_detected": False,
        "extraction_warnings": []
    }
    
    if not HAVE_FITZ:
        meta["extraction_warnings"].append("PyMuPDF (fitz) not installed. PDF extraction failed.")
        return "", meta
        
    text = ""
    try:
        doc = fitz.open(filepath)
        meta["pages"] = len(doc)
        
        for page in doc:
            page_text = page.get_text("text")
            text += page_text + "\n"
            
        doc.close()
    except Exception as e:
        meta["extraction_warnings"].append(f"PDF extraction error: {str(e)}")
        
    meta["character_count"] = len(text)
    
    if meta["pages"] > 0 and len(text.strip()) < meta["pages"] * 50:
        meta["suspected_scanned_pdf"] = True
        meta["ocr_required"] = True
        meta["extraction_warnings"].append("Very little text extracted. OCR may be required.")
        
    return text, meta

def extract_docx(filepath: str) -> tuple[str, dict]:
    meta = {
        "file_type": "docx",
        "extraction_method": "python-docx",
        "paragraph_count": 0,
        "character_count": 0,
        "suspected_scanned_pdf": False,
        "ocr_required": False,
        "legacy_doc_conversion_required": False,
        "table_structure_detected": False,
        "extraction_warnings": []
    }
    
    if not HAVE_DOCX:
        meta["extraction_warnings"].append("python-docx not installed. DOCX extraction failed.")
        return "", meta
        
    text = ""
    try:
        doc = docx.Document(filepath)
        
        # Check for tables
        if len(doc.tables) > 0:
            meta["table_structure_detected"] = True
            
        # We need reading order. doc.paragraphs does not include tables.
        # python-docx has doc.element.body which contains all elements in order.
        # But for simplicity, we can extract text from paragraphs and tables. 
        # A more robust way to get all text in order:
        for block in doc.iter_inner_content() if hasattr(doc, 'iter_inner_content') else doc.paragraphs:
            # We'll stick to a simpler approach if iter_inner_content is not available
            pass
            
        # Since python-docx iterators are complex, let's do paragraphs then tables as fallback, 
        # or better: use element body traversal.
        
        from docx.document import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
        
        def iter_block_items(parent):
            if isinstance(parent, Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                return
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)
                    
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text += block.text + "\n"
                meta["paragraph_count"] += 1
            elif isinstance(block, Table):
                for row in block.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    text += " | ".join(row_data) + "\n"
                    meta["paragraph_count"] += 1
                    
    except Exception as e:
        meta["extraction_warnings"].append(f"DOCX extraction error: {str(e)}")
        
    meta["character_count"] = len(text)
    return text, meta

def extract_text(filepath: str) -> tuple[str, dict]:
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext == ".docx":
        return extract_docx(filepath)
    elif ext == ".doc":
        meta = {
            "file_type": "doc",
            "extraction_method": "none",
            "legacy_doc_conversion_required": True,
            "character_count": 0,
            "extraction_warnings": ["Legacy .doc file detected. Local converter required."]
        }
        return "", meta
    elif ext in [".txt", ".md"]:
        text = ""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="latin-1") as f:
                text = f.read()
        meta = {
            "file_type": ext[1:],
            "extraction_method": "direct_read",
            "character_count": len(text)
        }
        return text, meta
    else:
        return "", {"extraction_warnings": [f"Unsupported file format: {ext}"]}

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", required=True)
    parser.add_argument("--output-dir", required=True)
    args = parser.parse_args()
    
    os.makedirs(args.output_dir, exist_ok=True)
    
    metadata_records = []
    
    for filename in os.listdir(args.input_dir):
        filepath = os.path.join(args.input_dir, filename)
        if not os.path.isfile(filepath):
            continue
            
        print(f"Extracting: {filename}")
        text, meta = extract_text(filepath)
        
        meta["filename"] = filename
        
        if text.strip():
            out_filename = os.path.splitext(filename)[0] + ".txt"
            out_filepath = os.path.join(args.output_dir, out_filename)
            with open(out_filepath, "w", encoding="utf-8") as f:
                f.write(text)
                
        metadata_records.append(meta)
        
    with open(os.path.join(args.output_dir, "extraction_metadata.json"), "w", encoding="utf-8") as f:
        json.dump(metadata_records, f, indent=2)
        
if __name__ == "__main__":
    main()
