"""
Standardize the DS05 mindfulness human focus group transcript.

Offline only. No API calls. Writes to
data/datasets_transcripts/standardized/mindfulness/fg1/, mirroring the
artefact set and turn schema already established for
data/datasets_transcripts/standardized/macho_meals/fg*/.

Nothing is invented. Two source conditions are recorded rather than repaired:

  1. An isolated EDITORIAL PREAMBLE is present in the source .docx: a paragraph
     reading "Here's the transcript with the spelling mistakes corrected:" sits
     mid-dialogue. The researcher confirmed (2026-08-04) that the transcript
     preserves the original speech and that this line did not form part of the
     focus-group dialogue. It is excluded from the analytical turn sequence and
     the removal is recorded, hash-anchored, in editorial_removal_record.json by
     scripts/mindfulness_editorial_removal_record.py.

     An earlier version of this file framed the line as evidence that the
     dialogue had been rewritten by a language model. THAT CLAIM IS WITHDRAWN.

  2. Transcript speaker labels ("Speaker 2".."Speaker 6") cannot be mapped to
     the five agent payloads in agents/mindfulness/ (MF_P1..MF_P5). No evidence
     for any particular assignment exists in either artefact. The counts match
     (5 and 5) and that is all that is recorded; identity_reconciliation.json
     reports matched=false for every speaker.

Usage:
    py scripts/standardize_mindfulness_baseline.py
"""

from __future__ import annotations

import hashlib
import json
import re
import sys
from pathlib import Path

from docx import Document

_ROOT = Path(__file__).resolve().parent.parent
_SOURCE = _ROOT / "data/datasets_transcripts/Mindfulness_raw transcript/Mindfulness_Focus Group Transcript.docx"
_OUT = _ROOT / "data/datasets_transcripts/standardized/mindfulness/fg1"

_SPEAKER_RE = re.compile(r"^([A-Za-z][A-Za-z0-9 ]{0,25}):\s*(.*)$", re.DOTALL)

_MODERATOR_LABEL = "Moderator"
_UNKNOWN_LABEL = "Unknown Speaker"
_UNATTRIBUTED_ID = "UNATTRIBUTED_PARTICIPANT"

# Non-dialogue lines to exclude from the turn sequence, recorded verbatim in
# baseline_metadata so the exclusion stays auditable.
_FRONT_MATTER_EXACT = {"Transcript"}
_EDITORIAL_PREAMBLE_MARKERS = (
    "here's the transcript with the spelling mistakes corrected",
)
_EDITORIAL_PREAMBLE_STATUS = (
    "EDITORIAL_PREAMBLE_REMOVED — RESEARCHER_CONFIRMED_NOT_PART_OF_SPEECH"
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _speaker_id(label: str) -> str:
    if label == _MODERATOR_LABEL:
        return "MODERATOR"
    if label == _UNKNOWN_LABEL:
        return _UNATTRIBUTED_ID
    slug = re.sub(r"[^a-z0-9]+", "_", label.lower()).strip("_")
    return f"mindfulness_fg1_{slug}"


def _role(label: str) -> str:
    if label == _MODERATOR_LABEL:
        return "moderator"
    if label == _UNKNOWN_LABEL:
        return "unattributed"
    return "participant"


def parse() -> dict:
    doc = Document(_SOURCE)
    paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]
    paragraphs = [(i, t) for i, t in paragraphs if t]

    turns: list[dict] = []
    front_matter: list[str] = []
    excluded: list[dict] = []
    editorial_preamble_paragraphs: list[dict] = []
    seen_dialogue = False

    for idx, text in paragraphs:
        low = text.lower()

        if any(marker in low for marker in _EDITORIAL_PREAMBLE_MARKERS):
            record = {"paragraph_index": idx, "text": text,
                      "reason": "editorial_preamble", "status": _EDITORIAL_PREAMBLE_STATUS}
            editorial_preamble_paragraphs.append(record)
            excluded.append(record)
            continue

        if not seen_dialogue and text in _FRONT_MATTER_EXACT:
            front_matter.append(text)
            continue

        match = _SPEAKER_RE.match(text)
        if match and _role(match.group(1)) in {"moderator", "participant", "unattributed"} and (
            match.group(1) == _MODERATOR_LABEL
            or match.group(1) == _UNKNOWN_LABEL
            or re.fullmatch(r"Speaker \d+", match.group(1))
        ):
            label, content = match.group(1), match.group(2).strip()
            seen_dialogue = True
            turns.append(
                {
                    "turn": len(turns),
                    "speaker_id": _speaker_id(label),
                    "canonical_speaker_id": _speaker_id(label),
                    "speaker_name": label,
                    "speaker_role": _role(label),
                    "content": content,
                    "source_type": "dataset_transcript",
                    "source_file": _SOURCE.name,
                    "original_file_type": "docx",
                    "page": None,
                    "paragraph_indices": [idx],
                    "standardization_confidence": "high",
                    "requires_review": False,
                }
            )
            continue

        if seen_dialogue and turns:
            # Unlabelled paragraph following a labelled turn: a continuation of
            # that same speaker's turn. Appended, with the paragraph index kept.
            turns[-1]["content"] = (turns[-1]["content"] + "\n\n" + text).strip()
            turns[-1]["paragraph_indices"].append(idx)
            continue

        if not seen_dialogue:
            front_matter.append(text)
            continue

        excluded.append({"paragraph_index": idx, "text": text, "reason": "unclassified"})

    return {
        "turns": turns,
        "front_matter": front_matter,
        "excluded": excluded,
        "editorial_preamble_paragraphs": editorial_preamble_paragraphs,
    }


def build_participant_metadata(turns: list[dict]) -> list[dict]:
    order: list[str] = []
    counts: dict[str, int] = {}
    names: dict[str, str] = {}
    roles: dict[str, str] = {}
    for t in turns:
        sid = t["speaker_id"]
        if sid not in counts:
            order.append(sid)
            counts[sid] = 0
            names[sid] = t["speaker_name"]
            roles[sid] = t["speaker_role"]
        counts[sid] += 1

    # Moderator first, then participants in first-appearance order, then unattributed.
    def sort_key(sid: str) -> tuple[int, int]:
        rank = {"moderator": 0, "participant": 1, "unattributed": 2}[roles[sid]]
        return (rank, order.index(sid))

    return [
        {
            "speaker_id": sid,
            "speaker_name": names[sid],
            "speaker_role": roles[sid],
            "canonical_speaker_id": sid,
            "turn_count": counts[sid],
        }
        for sid in sorted(order, key=sort_key)
    ]


def render_transcript(turns: list[dict]) -> str:
    return "\n\n".join(f"{t['speaker_name']}: {t['content']}" for t in turns) + "\n"


def main() -> int:
    if not _SOURCE.exists():
        print(f"FAIL: source not found: {_SOURCE}")
        return 2

    parsed = parse()
    turns = parsed["turns"]
    participants = build_participant_metadata(turns)

    participant_rows = [p for p in participants if p["speaker_role"] == "participant"]
    moderator_rows = [p for p in participants if p["speaker_role"] == "moderator"]
    unattributed_rows = [p for p in participants if p["speaker_role"] == "unattributed"]

    agent_dir = _ROOT / "agents/mindfulness"
    agent_ids = sorted(p.stem for p in agent_dir.glob("mf_*.json"))

    identity = [
        {
            "transcript_speaker": p["speaker_name"],
            "agent_id": None,
            "matched": False,
            "reason": (
                "No evidence in either the transcript or the agent payloads supports a "
                "mapping between numbered transcript speakers and MF_P1..MF_P5. Counts "
                "match (5 participants, 5 agents); assignment does not follow from that."
            ),
        }
        for p in participant_rows
    ]

    metadata = {
        "baseline_id": "mindfulness_fg1",
        "dataset": "DS05_SAM_MINDFULNESS",
        "focus_group": "FG1",
        "source_file": _SOURCE.name,
        "source_file_sha256": _sha256(_SOURCE),
        "original_file_type": "docx",
        "total_turns": len(turns),
        "moderator_turns": sum(p["turn_count"] for p in moderator_rows),
        "participant_turns": sum(p["turn_count"] for p in participant_rows),
        "unattributed_turns": sum(p["turn_count"] for p in unattributed_rows),
        "participant_count": len(participant_rows),
        "agents_available_count": len(agent_ids),
        "agents_matched": False,
        "agents_matched_count": 0,
        "no_matched_agents": True,
        "identity_mapping_status": "NOT_ESTABLISHED_NO_EVIDENCE",
        "researcher_edits_applied": None,
        "researcher_edit_note": (
            "Unknown. No edit log accompanies this source file, unlike the Macho Meals "
            "baselines whose researcher edits are documented."
        ),
        "SOURCE_INTEGRITY_FLAGS": {
            "editorial_preamble_removed": bool(parsed["editorial_preamble_paragraphs"]),
            "editorial_preamble_paragraphs": parsed["editorial_preamble_paragraphs"],
            "editorial_preamble_status": _EDITORIAL_PREAMBLE_STATUS,
            "removal_record": "editorial_removal_record.json",
            "researcher_confirmation": (
                "The researcher confirmed that the transcript preserves the original speech. "
                "One isolated editorial preamble was removed from the analytical copy and did "
                "not form part of the focus-group dialogue."
            ),
            "RETRACTED_EARLIER_CLAIM": (
                "A previous version of this metadata asserted that the transcript had been "
                "partially rewritten by a language model and that lexical/stylometric measures "
                "were therefore uninterpretable. That claim is WITHDRAWN. Lexical measures are "
                "admissible on this baseline, subject to the feasibility preflight in "
                "scripts/lexical_transportability_mindfulness.py."
            ),
            "identifiable_names_present_in_dialogue": (
                "The transcript contains first names spoken in dialogue (e.g. the moderator is "
                "addressed by name). This baseline is not fully de-identified."
            ),
        },
        "excluded_paragraphs": parsed["excluded"],
        "standardization_script": "scripts/standardize_mindfulness_baseline.py",
    }

    _OUT.mkdir(parents=True, exist_ok=True)
    (_OUT / "transcript.json").write_text(
        json.dumps(turns, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    (_OUT / "participant_metadata.json").write_text(
        json.dumps(participants, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    (_OUT / "identity_reconciliation.json").write_text(
        json.dumps(identity, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    (_OUT / "baseline_metadata.json").write_text(
        json.dumps(metadata, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    clean = render_transcript(turns)
    (_OUT / "clean_transcript.txt").write_text(clean, encoding="utf-8")
    (_OUT / "transcript.txt").write_text(clean, encoding="utf-8")
    raw_paras = [p.text.strip() for p in Document(_SOURCE).paragraphs if p.text.strip()]
    (_OUT / "raw_extracted_transcript.txt").write_text(
        "\n\n".join(raw_paras) + "\n", encoding="utf-8"
    )
    (_OUT / "front_matter.txt").write_text(
        "\n".join(parsed["front_matter"]) + "\n" if parsed["front_matter"] else "", encoding="utf-8"
    )
    (_OUT / "back_matter.txt").write_text("", encoding="utf-8")

    print(f"wrote {_OUT.relative_to(_ROOT)}")
    print(f"  turns:              {len(turns)}")
    print(f"  moderator turns:    {metadata['moderator_turns']}")
    print(f"  participant turns:  {metadata['participant_turns']}")
    print(f"  unattributed turns: {metadata['unattributed_turns']}")
    print(f"  participants:       {metadata['participant_count']} ({', '.join(p['speaker_name'] for p in participant_rows)})")
    print(f"  agents available:   {metadata['agents_available_count']} (mapping NOT established)")
    print(f"  excluded paragraphs: {len(parsed['excluded'])}")
    for rec in parsed["editorial_preamble_paragraphs"]:
        print(f"  editorial preamble removed at paragraph {rec['paragraph_index']}: {rec['text']!r}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
